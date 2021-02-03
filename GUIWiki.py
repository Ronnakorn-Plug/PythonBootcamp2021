# GUIWiki.py
import wikipedia

# python to docx
from docx import Document
def Wiki(keyword,lang='th'):
    wikipedia.set_lang(lang)

    # summary สำหรับบทความที่สรุป
    data = wikipedia.summary(keyword)

    # page + content บทความทั้งหน้า
    data2 = wikipedia.page(keyword)
    data2 = data2.content

    doc = Document() # สร้างfile word ใน python
    doc.add_heading(keyword,0)

    doc.add_paragraph(data2)
    doc.save(keyword + '.docx')
    print('สร้างไฟล์สำเร็จ')

# --Set ภาษาเป็น ไทย ------
wikipedia.set_lang('th')
from tkinter import *
from tkinter import ttk
from tkinter import messagebox # การสร้าง windown ขึ้นมาแสดงผลลัพท์

GUI = Tk()
GUI.geometry('400x300')
GUI.title('Wikipedia by Plug')

#---- Config --------
FONT1 = ('Angsana New',15)

# ------Label -------
L1 = ttk.Label(GUI, text=('ค้นหาบทความ'),font=FONT1,foreground='red')
L1.pack(pady=10)

# -----ช่องค้นหาข้อมูล ------
V_search = StringVar()
E1 = ttk.Entry(GUI,textvariable = V_search,font=FONT1,width=40)
E1.pack(pady=10)

# ----- Button -------
def Search():
    keyword = V_search.get()
    try: # try/except คือการลองค้นหาผลลัพท์ หรือ เรียกว่าการดักจับ error
        language = v_radio.get() # th / en / zh
        Wiki(keyword,language)
        messagebox.showinfo('บันทึกสำเร็จ','การค้นหาสำเร็จ บันทึกข้อมูลแล้ว') 
    except: # หากรันแล้วมีปัญหาจะแสดงการการติดทันที
        messagebox.showwarning('Keyword Error','กรุณากรอกคำค้นหาใหม่')
        
        
    # print(wikipedia.search(keyword)) # เป็นการกรองออกมาให้เราเลือกคำๆนั้นออกมาว่าเราต้องการคำไหน
    # result = wikipedia.summary(keyword)
    # print(result)
    
    

B1 = ttk.Button(GUI, text='Search',command=Search)
B1.pack(ipadx=20, ipady=10)

# เมนูเลือกภาษา
F1 = Frame(GUI)
F1.pack(pady=10)

v_radio = StringVar()

# Radiobutton คือการกำหนดเลือกช้อมูลได้อย่างเดียว 
RB1 = ttk.Radiobutton(F1,text='ภาษาไทย',variable=v_radio,value='th') 
RB2 = ttk.Radiobutton(F1,text='ภาษาอังกฤษ',variable=v_radio,value='en')
RB3 = ttk.Radiobutton(F1,text='ภาษาจีน',variable=v_radio,value='zh')
RB1.invoke() # สั่งให้ค่าเริ่มต้นเป็นภาษาไทย

RB1.grid(row=0,column=0) # grid คือกำหนดจุดวางข้อมูลเป็นแนวนอน pack คือแนวตั้ง
RB2.grid(row=0,column=1)
RB3.grid(row=0,column=2)



GUI.mainloop()
