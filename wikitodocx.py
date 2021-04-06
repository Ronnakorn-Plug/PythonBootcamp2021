# wikitodocx.py

import wikipedia
from docx import Document

wikipedia.set_lang('th')
product = ['ทุเรียน','มังคุด','องุ่น']

document = Document()
document.add_heading('รายการผลไม้', 0)

for pd in product:
	# content = wikipedia.summary(pd) # สรุปหลายรายการ
	content = wikipedia.page(pd).content
	document.add_heading(pd, level=1)
	p = document.add_paragraph('')
	p.add_run(content + '\n\n\n') # \n คือการ enter
	# document.add_page_break()


document.save('fruit.docx')